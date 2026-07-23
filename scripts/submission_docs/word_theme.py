from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image


BLUE = "2E74B5"
NAVY = "17324D"
ORANGE = "F28C28"
LIGHT = "F2F4F7"
GRAY = "5B6573"


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Cm(1.65)
    section.bottom_margin = Cm(1.55)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for style_name, size, color in (
        ("Title", 26, NAVY),
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, NAVY),
        ("Heading 3", 12, NAVY),
    ):
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)

    _add_page_footer(section)


def add_page_title(document: Document, number: str, title: str, subtitle: str | None = None) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(number.upper())
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(ORANGE)
    heading = document.add_heading(title, level=1)
    heading.paragraph_format.space_after = Pt(7)
    if subtitle:
        sub = document.add_paragraph(subtitle)
        sub.paragraph_format.space_after = Pt(9)
        sub_run = sub.runs[0]
        sub_run.italic = True
        sub_run.font.color.rgb = RGBColor.from_string(GRAY)


def add_body(document: Document, text: str, bold_prefix: str | None = None) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_prefix and text.startswith(bold_prefix):
        prefix = paragraph.add_run(bold_prefix)
        prefix.bold = True
        paragraph.add_run(text[len(bold_prefix) :])
    else:
        paragraph.add_run(text)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.add_run(item)


def add_numbered(document: Document, items: list[str]) -> None:
    for index, item in enumerate(items, start=1):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(4)
        number = paragraph.add_run(f"{index}.  ")
        number.bold = True
        paragraph.add_run(item)


def add_table(
    document: Document,
    headers: list[str],
    rows: list[list[object]],
    widths: list[float] | None = None,
    font_size: float = 8.5,
) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _shade(cell, NAVY)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(font_size)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for column_index, value in enumerate(values):
            cells[column_index].text = str(value)
            cells[column_index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index % 2:
                _shade(cells[column_index], LIGHT)
            for paragraph in cells[column_index].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Inches(width)


def add_picture_fit(
    document: Document,
    image_path: Path,
    max_width_inches: float = 6.7,
    max_height_inches: float = 8.6,
    caption: str | None = None,
) -> None:
    with Image.open(image_path) as image:
        width_px, height_px = image.size
    ratio = width_px / height_px
    width = max_width_inches
    height = width / ratio
    if height > max_height_inches:
        height = max_height_inches
        width = height * ratio
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image_path), width=Inches(width), height=Inches(height))
    if caption:
        caption_paragraph = document.add_paragraph(caption)
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_paragraph.paragraph_format.space_after = Pt(0)
        run = caption_paragraph.runs[0]
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(GRAY)


def add_callout(document: Document, title: str, body: str, color: str = BLUE) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.7)
    _shade(cell, LIGHT)
    _cell_border(cell, color)
    paragraph = cell.paragraphs[0]
    title_run = paragraph.add_run(f"{title}\n")
    title_run.bold = True
    title_run.font.color.rgb = RGBColor.from_string(color)
    body_run = paragraph.add_run(body)
    body_run.font.size = Pt(10.5)


def add_page_break(document: Document) -> None:
    document.add_page_break()


def _add_page_footer(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("NHÓM 01  •  LAB211 FLASH SALE  •  TRANG ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(GRAY)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def _shade(cell, color: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    fill = properties.find(qn("w:shd"))
    if fill is None:
        fill = OxmlElement("w:shd")
        properties.append(fill)
    fill.set(qn("w:fill"), color)


def _cell_border(cell, color: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    borders = properties.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        properties.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "8" if edge == "left" else "2")
        tag.set(qn("w:color"), color)
        borders.append(tag)
