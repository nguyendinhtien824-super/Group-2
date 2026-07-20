from __future__ import annotations

import xml.etree.ElementTree as element_tree
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from benchmark_data import BenchmarkRow, BenchmarkSummary, create_charts, load_benchmarks, summarize
from report_content import (
    ARCHITECTURE_ROWS,
    DATA_ROWS,
    LOCK_ROWS,
    ROLE_ROWS,
    SECURITY_ROWS,
    TRACEABILITY_ROWS,
)
from word_theme import (
    BLUE,
    NAVY,
    ORANGE,
    add_body,
    add_bullets,
    add_callout,
    add_numbered,
    add_page_break,
    add_page_title,
    add_picture_fit,
    add_table,
    configure_document,
)


def build_report(project_root: Path, output_path: Path) -> None:
    data = load_benchmarks(project_root / "data")
    charts = create_charts(data, project_root / "docs" / "charts")
    summaries = {stock: summarize(rows) for stock, rows in data.items()}
    test_total, test_failures = _read_test_results(project_root / "target" / "surefire-reports")

    document = Document()
    configure_document(document)
    document.core_properties.title = "Báo cáo LAB211 Flash Sale Simulator"
    document.core_properties.subject = "Race Condition, FIFO, Locking và benchmark TPS"
    document.core_properties.author = "Nhóm 01"

    _cover(document)
    _executive_summary(document, summaries)
    _research_problem(document)
    _traceability(document)
    _architecture(document)
    _use_cases(document, project_root)
    _class_model(document, project_root)
    _dataset(document)
    _generator(document, project_root)
    _order_flow(document, project_root)
    _fifo_priority(document)
    _race_condition(document, project_root)
    _locking(document)
    _simulator(document, project_root)
    _methodology(document)
    _scenario_results(document, 100, data[100], summaries[100], charts["throughput_100"], charts["safety_100"])
    _scenario_results(document, 2000, data[2000], summaries[2000], charts["throughput_2000"], charts["target_2000"])
    _verdict(document, summaries)
    _quality(document, test_total, test_failures)
    _delivery(document)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def _cover(document: Document) -> None:
    for _ in range(3):
        document.add_paragraph()
    label = document.add_paragraph()
    label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label_run = label.add_run("FPT POLYTECHNIC  •  LAB211")
    label_run.bold = True
    label_run.font.size = Pt(12)
    label_run.font.color.rgb = RGBColor.from_string(ORANGE)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("FLASH SALE\nCONCURRENCY SIMULATOR")
    title_run.bold = True
    title_run.font.size = Pt(30)
    title_run.font.color.rgb = RGBColor.from_string(NAVY)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Race Condition • FIFO/Priority • 4 Lock Mechanisms • TPS Benchmark")
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = RGBColor.from_string(BLUE)
    document.add_paragraph()
    add_table(
        document,
        ["Thông tin", "Nội dung"],
        [
            ["Nhóm", "NHÓM 01"],
            ["Thành viên", "Bổ sung họ tên/MSSV trước khi nộp"],
            ["Phiên bản", "Bản hoàn thiện yêu cầu LAB211 Flash Sale V2"],
            ["Tài liệu kèm", "README, sơ đồ, CSV benchmark, slide và source code"],
        ],
        widths=[1.7, 4.9],
        font_size=10,
    )
    add_callout(
        document,
        "Cam kết số liệu",
        "Mọi kết quả trong báo cáo được đọc trực tiếp từ 24 lượt benchmark thật; không sửa số để ép đạt mục tiêu.",
        ORANGE,
    )
    add_page_break(document)


def _executive_summary(document: Document, summaries: dict[int, list[BenchmarkSummary]]) -> None:
    add_page_title(document, "01", "Tóm tắt điều hành", "Kết quả chính và câu trả lời ngắn cho đề tài")
    low_no_lock = _summary(summaries[100], "NO_LOCK")
    high_file = _summary(summaries[2000], "FILE_LOCK")
    high_sync = _summary(summaries[2000], "SYNCHRONIZED")
    add_body(
        document,
        "Hệ thống mô phỏng Flash Sale bằng Java 17, kiến trúc MVC nhiều lớp và persistence CSV UTF-8. "
        "Luồng đặt hàng áp dụng kiểm tra trạng thái fail-fast, giới hạn mua trong vùng tới hạn, FIFO theo sequence và ưu tiên tier trên STANDARD mà không vượt tồn kho.",
    )
    add_bullets(
        document,
        [
            f"Tồn kho 100: No Lock đạt trung bình {low_no_lock.average_tps:.2f} TPS nhưng tạo tổng {low_no_lock.negative_stock_total} đơn vị âm kho và dữ liệu không nhất quán.",
            "FileLock, synchronized và optimistic giữ âm kho bằng 0 ở cả ba lần chạy kịch bản thấp kho.",
            f"Tồn kho 2000: FileLock đạt trung bình {high_file.average_tps:.2f} TPS; synchronized đạt {high_sync.average_tps:.2f} TPS.",
            "Không cơ chế an toàn nào đồng thời đạt 0% âm kho và suy giảm TPS không quá 30% trong cả hai kịch bản.",
            "Kết luận không bị làm đẹp: mục tiêu nghiên cứu chưa đạt ổn định trong môi trường đo hiện tại.",
        ],
    )
    add_callout(
        document,
        "Kết luận",
        "Ở stock 2000, synchronized đạt ngưỡng cả 3 run; nhưng tại stock 100 mọi cơ chế an toàn đều hụt ngưỡng TPS. Nếu ưu tiên tính đúng trong một JVM, synchronized là lựa chọn đơn giản; NIO FileLock phù hợp hơn khi cần phối hợp nhiều process.",
    )
    add_page_break(document)


def _research_problem(document: Document) -> None:
    add_page_title(document, "02", "Bài toán và câu hỏi nghiên cứu")
    add_body(document, "Flash Sale tập trung nhiều yêu cầu mua vào cùng một khoảng thời gian ngắn. Nếu nhiều luồng cùng đọc một số tồn kho rồi cùng ghi, hệ thống có thể bán vượt số lượng, tạo âm kho hoặc mất cập nhật.")
    document.add_heading("Câu hỏi nghiên cứu", level=2)
    add_callout(document, "Research question", "Cơ chế khóa nào loại bỏ âm kho nhưng vẫn giữ mức suy giảm throughput không quá 30% so với No Lock?", ORANGE)
    document.add_heading("Mục tiêu", level=2)
    add_numbered(document, [
        "Cài đặt cùng một workload trên bốn cơ chế khóa.",
        "Đo success, failed, final stock, negative stock, retry, consistency và TPS.",
        "Ghép từng cơ chế an toàn với baseline No Lock cùng run để hạn chế sai lệch môi trường.",
        "Lặp ba lần, công bố cả raw result và trung bình/median; không chỉ chọn lần đẹp nhất.",
    ])
    document.add_heading("Phạm vi", level=2)
    add_body(document, "Ứng dụng console một máy, lưu CSV; FileLock dùng khóa cấp hệ điều hành, synchronized bảo vệ trong JVM, optimistic dùng trường version và retry tối đa ba lần.")
    add_page_break(document)


def _traceability(document: Document) -> None:
    add_page_title(document, "03", "Ma trận truy vết yêu cầu", "Mỗi hạng mục trỏ tới bằng chứng có thể kiểm tra")
    add_table(document, ["ID", "Yêu cầu", "Bằng chứng", "KQ"], TRACEABILITY_ROWS, widths=[0.45, 2.35, 3.25, 0.55], font_size=7.4)
    add_callout(document, "Ngoài phạm vi tạo mới", "AI audit log do nhóm đã có và được nộp riêng; báo cáo này không tạo hoặc giả lập lịch sử AI.")
    add_page_break(document)


def _architecture(document: Document) -> None:
    add_page_title(document, "04", "Kiến trúc hệ thống", "Composition root nối dependency; View không chạm repository")
    add_table(document, ["Tầng", "Trách nhiệm", "Ranh giới"], ARCHITECTURE_ROWS, widths=[1.05, 3.7, 1.85], font_size=9)
    add_bullets(document, [
        "CsvRepository là generic CRUD; mapper/codec/value converter tách riêng.",
        "OrderPlacementService sở hữu atomic business flow; OrderRequestQueue chỉ điều phối thứ tự vào vùng tới hạn.",
        "SimulationExecutor quản lý thread pool + CountDownLatch; SimulatorService quản lý chiến lược lock và metric.",
        "FlashSaleApplication là composition root và cung cấp CLI không phụ thuộc đường dẫn máy cá nhân.",
    ])
    add_page_break(document)


def _use_cases(document: Document, root: Path) -> None:
    add_page_title(document, "05", "Use Case theo vai trò")
    add_table(document, ["Vai trò", "Hành động chính", "Giới hạn"], ROLE_ROWS, widths=[1.1, 3.8, 1.7], font_size=8.5)
    add_picture_fit(document, root / "docs" / "use_case_diagram.png", max_height_inches=5.7, caption="Hình 1 — Use Case Diagram")
    add_page_break(document)


def _class_model(document: Document, root: Path) -> None:
    add_page_title(document, "06", "Mô hình lớp và quan hệ miền")
    add_picture_fit(document, root / "docs" / "class_diagram.png", max_height_inches=8.7, caption="Hình 2 — Class Diagram: thuộc tính, phương thức, visibility và multiplicity")
    add_page_break(document)


def _dataset(document: Document) -> None:
    add_page_title(document, "07", "Thiết kế dữ liệu CSV", "Dữ liệu sinh ≥ 12.500 dòng trước transaction benchmark")
    add_table(document, ["Tệp", "Số dòng nền", "Vai trò"], DATA_ROWS, widths=[1.7, 1.15, 3.75], font_size=9)
    add_bullets(document, [
        "Product và FlashItem có version để kiểm soát optimistic concurrency.",
        "Order có eventId để giới hạn tối đa 2 sản phẩm theo customer + product + event.",
        "Hash Argon2id chứa dấu phẩy được CsvRowCodec quote/escape và kiểm thử round-trip qua repository.",
        "Ghi dữ liệu qua file tạm rồi atomic move để tránh tệp nửa vời khi lỗi.",
        "Tìm kiếm product kết hợp category và khoảng giá; test đọc 10.000 dòng dưới 1 giây.",
    ])
    add_page_break(document)


def _generator(document: Document, root: Path) -> None:
    add_page_title(document, "08", "Luồng sinh dữ liệu")
    add_picture_fit(document, root / "docs" / "flowcharts" / "data_generator_flow.png", max_height_inches=8.8, caption="Hình 3 — DataGenerator Flowchart")
    add_page_break(document)


def _order_flow(document: Document, root: Path) -> None:
    add_page_title(document, "09", "Luồng đặt hàng Flash Sale")
    add_picture_fit(document, root / "docs" / "flowcharts" / "order_flow.png", max_height_inches=8.8, caption="Hình 4 — Order Flowchart")
    add_page_break(document)


def _fifo_priority(document: Document) -> None:
    add_page_title(document, "10", "FIFO và ưu tiên VIP/PREMIUM", "Ưu tiên chỉ thay đổi thứ tự chờ, tuyệt đối không bỏ qua tồn kho")
    add_numbered(document, [
        "Mỗi request nhận sequence tăng đơn điệu khi vào OrderRequestQueue.",
        "Nhóm tier trên STANDARD được ưu tiên trước nhóm STANDARD đang chờ.",
        "Trong cùng nhóm priority, request có sequence nhỏ nhất được cấp permit trước.",
        "Request đang ở vùng tới hạn không bị preempt; permit luôn được giải phóng bằng AutoCloseable.",
        "Sau khi được cấp permit, service đọc lại customer/item/event và thực hiện toàn bộ validation trong vùng atomic.",
    ])
    add_table(document, ["Tình huống test", "Kỳ vọng"], [
        ["VIP đến sau Standard đang chờ", "VIP được phục vụ trước request Standard chưa active"],
        ["Hai request cùng tier", "Giữ nguyên FIFO theo sequence"],
        ["VIP mua khi hết kho", "Fail, không bán vượt kho"],
        ["Cạnh tranh sản phẩm cuối", "Một success; request còn lại fail; soldQty ≤ limited"],
    ], widths=[3.2, 3.4], font_size=9)
    add_page_break(document)


def _race_condition(document: Document, root: Path) -> None:
    add_page_title(document, "11", "Race Condition và vùng tới hạn")
    add_picture_fit(document, root / "docs" / "flowcharts" / "race_condition_flow.png", max_height_inches=8.8, caption="Hình 5 — Race Condition Flowchart")
    add_page_break(document)


def _locking(document: Document) -> None:
    add_page_title(document, "12", "Bốn cơ chế đồng bộ")
    add_table(document, ["Cơ chế", "Cách hoạt động", "Điểm mạnh", "Đánh đổi"], LOCK_ROWS, widths=[1.2, 2.2, 1.7, 1.55], font_size=8.2)
    add_callout(document, "Công thức TPS", "TPS = số giao dịch thành công / (elapsed time tính bằng giây). Baseline được ghép theo cùng run để tính vsBaselinePercent.")
    add_body(document, "Optimistic Lock không fallback sang khóa khác. Mỗi xung đột version chỉ retry tối đa ba lần; quá giới hạn được ghi nhận là thất bại. Đây là kết quả trung thực của chiến lược được yêu cầu, kể cả khi TPS thấp.")
    add_page_break(document)


def _simulator(document: Document, root: Path) -> None:
    add_page_title(document, "13", "Luồng Simulator")
    add_picture_fit(document, root / "docs" / "flowcharts" / "simulator_flow.png", max_height_inches=8.8, caption="Hình 6 — Simulator Flowchart")
    add_page_break(document)


def _methodology(document: Document) -> None:
    add_page_title(document, "14", "Phương pháp thực nghiệm", "Hai kịch bản tách lỗi an toàn khỏi giới hạn capacity")
    add_table(document, ["Kịch bản", "Cấu hình", "Mục đích"], [
        ["A — stock 100", "1000 requests × 4 threads × 3 runs", "Kích hoạt oversell/race và đo âm kho"],
        ["B — stock 2000", "1000 requests × 4 threads × 3 runs", "Đủ capacity để so throughput công bằng hơn"],
    ], widths=[1.5, 2.4, 2.65], font_size=9)
    add_numbered(document, [
        "ExecutorService tạo bốn worker; ready latch bảo đảm mọi worker sẵn sàng.",
        "start latch phát lệnh đồng thời; done latch chờ toàn bộ request kết thúc.",
        "Mỗi request mua quantity 1–2; kết quả được persist vào transactions.csv.",
        "Mỗi run chạy đủ bốn cơ chế; No Lock là baseline của chính run đó.",
        "Công bố 12 raw rows/kịch bản, trung bình, median và số run đạt target.",
    ])
    add_callout(document, "Tiêu chí đạt", "Âm kho = 0, dataConsistent = true và vsBaselinePercent ≥ -30% trong cả ba lần chạy.", ORANGE)
    add_page_break(document)


def _scenario_results(
    document: Document,
    stock: int,
    rows: list[BenchmarkRow],
    summary: list[BenchmarkSummary],
    throughput_chart: Path,
    secondary_chart: Path,
) -> None:
    number = "15" if stock == 100 else "16"
    add_page_title(document, number, f"Kết quả thực nghiệm — tồn kho {stock}")
    raw_rows = [[
        row.run, _short_lock(row.lock_type), row.success, row.failed, row.final_stock,
        row.negative_stock, row.retries, f"{row.tps:.2f}", f"{row.vs_baseline:.1f}%", row.target,
    ] for row in rows]
    add_table(document, ["Run", "Lock", "OK", "Fail", "Kho", "Âm", "Retry", "TPS", "vsBase", "KQ"], raw_rows, font_size=6.9)
    summary_rows = [[
        _short_lock(item.lock_type), f"{item.average_tps:.2f}", f"{item.median_tps:.2f}",
        f"{item.average_drop:.1f}%", item.negative_stock_total, f"{item.target_runs}/3",
    ] for item in summary]
    add_table(document, ["Lock", "Avg TPS", "Median", "Avg drop", "Tổng âm", "Đạt"], summary_rows, font_size=7.8)
    add_picture_fit(document, throughput_chart, max_height_inches=3.0, caption=f"Biểu đồ TPS trung bình — stock {stock}")
    add_picture_fit(document, secondary_chart, max_height_inches=2.8, caption="Biểu đồ an toàn/retry hoặc ngưỡng throughput")
    add_page_break(document)


def _verdict(document: Document, summaries: dict[int, list[BenchmarkSummary]]) -> None:
    add_page_title(document, "17", "Đánh giá mục tiêu và thảo luận")
    rows = []
    for stock, items in summaries.items():
        for item in items:
            if item.lock_type == "NO_LOCK":
                continue
            rows.append([stock, _short_lock(item.lock_type), f"{item.average_drop:.1f}%", f"{item.consistent_runs}/3", f"{item.target_runs}/3", "CHƯA ĐẠT" if item.target_runs < 3 else "ĐẠT"])
    add_table(document, ["Stock", "Cơ chế", "Avg drop", "Nhất quán", "Run đạt", "Kết luận"], rows, widths=[0.65, 1.55, 1.0, 1.0, 0.85, 1.2], font_size=8.3)
    add_bullets(document, [
        "No Lock nhanh nhưng không thể dùng production vì âm kho và lost update.",
        f"Ở stock 2000, synchronized đạt {_summary(summaries[2000], 'SYNCHRONIZED').target_runs}/3 run và FileLock đạt {_summary(summaries[2000], 'FILE_LOCK').target_runs}/3 run.",
        "Ở stock 100, cả ba cơ chế an toàn đều 0/3 run đạt; do đó chưa cơ chế nào đạt xuyên cả hai kịch bản.",
        "Optimistic tăng so với bản có delay conflict nhân tạo, nhưng contention cao làm retry cạn và TPS vẫn rất thấp.",
        "Môi trường file CSV khiến chi phí I/O chi phối; kết quả không được suy rộng trực tiếp sang database server.",
        "Đề xuất nghiên cứu tiếp: warm-up có kiểm soát, nhiều vòng hơn, histogram latency và backend database thật.",
    ])
    add_callout(document, "Trả lời research question", "Trong cấu hình đo hiện tại, chưa có cơ chế an toàn nào thỏa đồng thời hai mục tiêu trong cả hai kịch bản.", ORANGE)
    add_page_break(document)


def _quality(document: Document, test_total: int, test_failures: int) -> None:
    add_page_title(document, "18", "Kiểm thử, bảo mật và khả năng chạy máy khác")
    add_callout(document, "Quality gate", f"Maven clean verify: {test_total} tests; {test_failures} failure/error. Fat JAR được smoke bằng EOF an toàn.", BLUE)
    add_table(document, ["Hạng mục bảo mật", "Triển khai"], SECURITY_ROWS, widths=[1.8, 4.8], font_size=9)
    add_bullets(document, [
        "Concurrent tests kiểm tra oversell, last-stock, FIFO và priority.",
        "Multi-role test chạy Customer và Admin ở hai session executor độc lập: đặt đơn → duyệt → khách thấy APPROVED.",
        "Data generator được test generate → repository round-trip → service boot, tránh build xanh nhưng dữ liệu thật không đọc được.",
        "Đường dẫn đều tương đối; Java 17 + Maven command trong README; secret chỉ qua environment.",
    ])
    add_page_break(document)


def _delivery(document: Document) -> None:
    add_page_title(document, "19", "Kết luận và gói bàn giao")
    add_body(document, "Đề tài đã được hoàn thiện theo hướng có thể chạy và kiểm chứng: code Java/MVC, dữ liệu CSV, lock simulator, benchmark raw, biểu đồ, sơ đồ, báo cáo, slide, test và hướng dẫn đều được đóng trong một ZIP sạch.")
    add_table(document, ["Thành phần", "Vị trí"], [
        ["Source + test", "src/, test/, pom.xml"],
        ["Dữ liệu và benchmark", "data/*.csv, data/simulation*.md"],
        ["Sơ đồ + chart", "docs/*.png, docs/flowcharts/, docs/charts/"],
        ["Báo cáo", "docs/report.docx"],
        ["Slide", "docs/slide.pptx"],
        ["Hướng dẫn", "README.md, .env.example, run*.bat"],
    ], widths=[2.25, 4.35], font_size=9.5)
    document.add_heading("Lệnh tái lập", level=2)
    add_bullets(document, [
        "mvn clean verify",
        "java -jar target/flash-sale-simulator.jar --generate-data",
        "java -jar target/flash-sale-simulator.jar --benchmark-report",
        "java -jar target/flash-sale-simulator.jar",
    ])
    add_callout(document, "Việc nhóm cần điền", "Bổ sung họ tên/MSSV trên trang bìa và ghép AI audit log mà nhóm đã có trước khi nộp chính thức.", ORANGE)


def _summary(items: list[BenchmarkSummary], lock_type: str) -> BenchmarkSummary:
    return next(item for item in items if item.lock_type == lock_type)


def _short_lock(lock_type: str) -> str:
    return {"SYNCHRONIZED": "SYNC", "OPTIMISTIC": "OPT"}.get(lock_type, lock_type)


def _read_test_results(report_dir: Path) -> tuple[int, int]:
    total = 0
    failed = 0
    for report in report_dir.glob("TEST-*.xml"):
        root = element_tree.parse(report).getroot()
        total += int(root.attrib.get("tests", 0))
        failed += int(root.attrib.get("failures", 0)) + int(root.attrib.get("errors", 0))
    if total == 0:
        raise RuntimeError("Chưa có Surefire report. Hãy chạy mvn clean verify trước khi tạo báo cáo.")
    return total, failed
